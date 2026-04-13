from __future__ import annotations

import os
import html
import base64
import io
from html.parser import HTMLParser
from pathlib import Path
from datetime import datetime
from typing import List, Optional

from PyQt5 import QtCore, QtGui, QtWidgets, QtPrintSupport

# Accessing classes directly to help some linters
Qt = QtCore.Qt
pyqtSignal = QtCore.pyqtSignal
QPoint = QtCore.QPoint
QByteArray = QtCore.QByteArray
QBuffer = QtCore.QBuffer
QIODevice = QtCore.QIODevice
QPrinter = QtPrintSupport.QPrinter
QBrush = QtGui.QBrush
QPen = QtGui.QPen
QFont = QtGui.QFont
QPixmap = QtGui.QPixmap
QTextCharFormat = QtGui.QTextCharFormat
QTextDocument = QtGui.QTextDocument
QPainter = QtGui.QPainter
QPen = QtGui.QPen
QColor = QtGui.QColor
QImage = QtGui.QImage
QTextTableFormat = QtGui.QTextTableFormat
QTextLength = QtGui.QTextLength
QTextCursor = QtGui.QTextCursor
QTextListFormat = QtGui.QTextListFormat
QAction = QtWidgets.QAction
QMenu = QtWidgets.QMenu
QComboBox = QtWidgets.QComboBox
QDialog = QtWidgets.QDialog
QDialogButtonBox = QtWidgets.QDialogButtonBox
QFormLayout = QtWidgets.QFormLayout
QHBoxLayout = QtWidgets.QHBoxLayout
QLabel = QtWidgets.QLabel
QListWidgetItem = QtWidgets.QListWidgetItem
QLineEdit = QtWidgets.QLineEdit
QMainWindow = QtWidgets.QMainWindow
QMessageBox = QtWidgets.QMessageBox
QPushButton = QtWidgets.QPushButton
QCheckBox = QtWidgets.QCheckBox
QSpinBox = QtWidgets.QSpinBox
QSizePolicy = QtWidgets.QSizePolicy
QTextEdit = QtWidgets.QTextEdit
QToolBar = QtWidgets.QToolBar
QVBoxLayout = QtWidgets.QVBoxLayout
QWidget = QtWidgets.QWidget
QFileDialog = QtWidgets.QFileDialog
QFontComboBox = QtWidgets.QFontComboBox

from models.document import Document, _load_ref_body_html
from models.user import User
from services.comment_service import CommentService
from services.workflow_service import WorkflowService
from ui.components.comment_widget import CommentWidget
from ui.components.signature_widget import SignatureWidget

import traceback
try:
    import win32com.client
except ImportError:
    win32com = None


def _build_ref_and_date(*, workflow: WorkflowService, doc: Document, user: User) -> tuple[str, str]:
    # Use the document creator instead of the current user to keep the REF and labels persistent.
    creator_user = workflow.get_user(doc.created_by) if doc.id and doc.created_by else user
    target_user = creator_user if creator_user else user

    created_at = (getattr(doc, "created_at", "") or "").strip()
    dt: datetime
    try:
        dt = datetime.fromisoformat(created_at) if created_at else datetime.now()
    except Exception:
        dt = datetime.now()
    
    date_str = dt.strftime("%d/%m/%Y")
    day = dt.strftime("%d")
    yy = dt.strftime("%y")
    iso_date = dt.strftime("%Y-%m-%d")
    
    seq = workflow.daily_sequence_for_document(document_id=getattr(doc, "id", None), iso_date=iso_date)
    
    # Use target_user (original creator) for dept/lab info
    dept = (getattr(target_user, "department", "") or "").strip() or "DEPT"
    lab = (getattr(target_user, "lab", "") or "").strip() or "LAB"
    dept = dept.replace("/", "-")
    lab = lab.replace("/", "-")
    ref_prefix = workflow.get_setting("ref_prefix", "NECOP/")
    if not ref_prefix.endswith("/"):
        ref_prefix += "/"

    ref_value = f"{ref_prefix}{dept}/{lab}/{day}{yy}-{day}{seq:02d}"
    return ref_value, date_str


def _is_signature_artifact_line(line: str) -> bool:
    s = (line or "").strip().lower()
    if not s:
        return False
    if s.startswith("approved by"):
        return True
    if s.startswith("timestamp:"):
        return True
    return False


class _HtmlBodyParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.items: List[dict] = []
        self._in_table = False
        self._in_tr = False
        self._in_td = False
        self._ignore = False
        self._cur_p: List[dict] = [] # List of {text, font, size, bold, italic, underline}
        self._cur_table: List[List[List[dict]]] = []
        self._cur_row: List[List[dict]] = []
        self._cur_cell: List[dict] = []
        self._in_list = False
        self._list_type = None
        self._cur_list_items: List[dict] = []
        self._list_indent_stack: List[float] = [0.0]
        
        # Style stack
        self._font_stack = ["Times New Roman"]
        self._size_stack = [12.0]
        self._bold_stack = [False]
        self._italic_stack = [False]
        self._under_stack = [False]

    def handle_starttag(self, tag: str, attrs) -> None:
        t = (tag or "").lower()
        if t in ("style", "script", "head", "meta", "link"):
            self._ignore = True
            return

        if t == "table":
            self._flush_p()
            self._in_table = True
            self._cur_table = []
            return
        
        if t in ("ul", "ol"):
            self._flush_p()
            self._in_list = True
            # Capture indentation from QTextEdit HTML, e.g.
            # style="margin-top: 0px; margin-bottom: 0px; margin-left: 0px; -qt-list-indent: 1;"
            style_attr = next((v for k, v in attrs if (k or "").lower() == "style"), "") or ""
            print(f"DEBUG: List style_attr = {style_attr}")
            indent_px: Optional[float] = None
            indent_level: Optional[int] = None
            if style_attr:
                try:
                    import re
                    m = re.search(r"margin-left\s*:\s*([0-9.]+)px", style_attr)
                    if m:
                        indent_px = float(m.group(1))
                    m2 = re.search(r"-qt-list-indent\s*:\s*([0-9]+)", style_attr)
                    if m2:
                        indent_level = int(m2.group(1))
                except Exception:
                    indent_px = None
                    indent_level = None

            # Prefer margin-left when available AND non-zero; otherwise use qt indent level.
            # Convert px->pt with ~0.75 factor for margin-left.
            # Each indent level is approximately 18pt (about 0.25 inch).
            if indent_px is not None and indent_px > 0:
                indent_pt = float(indent_px) * 0.75
            elif indent_level is not None:
                indent_pt = float(indent_level) * 18.0
            else:
                indent_pt = self._list_indent_stack[-1]
            print(f"DEBUG: Calculated indent_pt = {indent_pt}")
            self._list_indent_stack.append(indent_pt)

            if t == "ul":
                self._list_type = "bullet"
            else:
                # Check for list style type to determine numbering format
                style_str = next((v for k, v in attrs if k.lower() == "type"), "")
                if style_str.lower() == "a":
                    self._list_type = "alpha"
                elif style_str.lower() == "i":
                    self._list_type = "roman"
                else:
                    self._list_type = "number"
            self._cur_list_items = []
            return
        
        if t == "li" and self._in_list:
            self._cur_p = []
            return
        
        # Handle styles
        if t == "b" or t == "strong": self._bold_stack.append(True)
        elif t == "i" or t == "em": self._italic_stack.append(True)
        elif t == "u": self._under_stack.append(True)
        elif t == "span":
            style_str = next((v for k, v in attrs if k.lower() == "style"), "")
            font = self._font_stack[-1]
            size = self._size_stack[-1]
            if "font-family" in style_str:
                import re
                m = re.search(r"font-family\s*:\s*'([^']+)'", style_str)
                if not m: m = re.search(r"font-family\s*:\s*([^;]+)", style_str)
                if m: font = m.group(1).strip().strip("'").strip('"')
            if "font-size" in style_str:
                import re
                m = re.search(r"font-size\s*:\s*(\d+)pt", style_str)
                if m: size = float(m.group(1))
            self._font_stack.append(font)
            self._size_stack.append(size)
        else:
            # push defaults for tags that don't change style to keep stacks aligned
            if t not in ("br", "table", "tr", "td", "th"):
                self._bold_stack.append(self._bold_stack[-1])
                self._italic_stack.append(self._italic_stack[-1])
                self._under_stack.append(self._under_stack[-1])
                self._font_stack.append(self._font_stack[-1])
                self._size_stack.append(self._size_stack[-1])

        if self._in_table:
            if t == "tr":
                self._in_tr = True
                self._cur_row = []
                return
            if t in ("td", "th"):
                self._in_td = True
                self._cur_cell = []
                return
            if t == "br" and self._in_td:
                self._cur_cell.append({"text": "\n", "font": self._font_stack[-1], "size": self._size_stack[-1], 
                                     "bold": self._bold_stack[-1], "italic": self._italic_stack[-1], "underline": self._under_stack[-1]})
                return
            return

        if t == "br":
            self._cur_p.append({"text": "\n", "font": self._font_stack[-1], "size": self._size_stack[-1], 
                               "bold": self._bold_stack[-1], "italic": self._italic_stack[-1], "underline": self._under_stack[-1]})
            return

    def handle_endtag(self, tag: str) -> None:
        t = (tag or "").lower()
        if t in ("style", "script", "head", "meta", "link"):
            self._ignore = False
            return

        # Pop styles
        if t == "b" or t == "strong": 
            if len(self._bold_stack) > 1: self._bold_stack.pop()
        elif t == "i" or t == "em": 
            if len(self._italic_stack) > 1: self._italic_stack.pop()
        elif t == "u": 
            if len(self._under_stack) > 1: self._under_stack.pop()
        elif t not in ("br", "table", "tr", "td", "th", "p", "div"):
            if len(self._font_stack) > 1: self._font_stack.pop()
            if len(self._size_stack) > 1: self._size_stack.pop()
            if len(self._bold_stack) > 1: self._bold_stack.pop()
            if len(self._italic_stack) > 1: self._italic_stack.pop()
            if len(self._under_stack) > 1: self._under_stack.pop()

        if t == "table":
            if self._cur_table:
                self.items.append({"type": "table", "rows": self._cur_table})
            self._in_table = False
            self._in_tr = False
            self._in_td = False
            self._cur_table = []
            self._cur_row = []
            self._cur_cell = []
            return

        if self._in_table:
            if t in ("td", "th") and self._in_td:
                self._cur_row.append(list(self._cur_cell))
                self._in_td = False
                self._cur_cell = []
                return
            if t == "tr" and self._in_tr:
                self._cur_table.append(list(self._cur_row))
                self._in_tr = False
                self._cur_row = []
                return
            return

        if t in ("p", "div"):
            self._flush_p()
            return
        
        if t in ("ul", "ol") and self._in_list:
            if self._cur_list_items:
                self.items.append({
                    "type": "list",
                    "list_type": self._list_type,
                    "items": self._cur_list_items,
                    "indent_pt": float(self._list_indent_stack[-1] if self._list_indent_stack else 0.0),
                })
            self._in_list = False
            self._list_type = None
            self._cur_list_items = []
            if len(self._list_indent_stack) > 1:
                self._list_indent_stack.pop()
            return
        
        if t == "li" and self._in_list and self._cur_p:
            # Only keep non-empty list items (avoid whitespace-only items from HTML formatting)
            txt = "".join((c.get("text", "") or "") for c in self._cur_p)
            if txt.strip() != "":
                self._cur_list_items.append({"chunks": list(self._cur_p)})
            self._cur_p = []
            return

    def handle_data(self, data: str) -> None:
        if self._ignore or not data:
            return
        # QTextEdit HTML often contains formatting whitespace/newlines between tags.
        # Treat whitespace-only nodes as non-content so they don't become empty DOCX paragraphs.
        if data.strip() == "" and "\xa0" not in data:
            return
        chunk = {
            "text": data,
            "font": self._font_stack[-1],
            "size": self._size_stack[-1],
            "bold": self._bold_stack[-1],
            "italic": self._italic_stack[-1],
            "underline": self._under_stack[-1]
        }
        if self._in_table and self._in_td:
            self._cur_cell.append(chunk)
            return
        self._cur_p.append(chunk)

    def close(self) -> None:
        self._flush_p()
        super().close()

    def _flush_p(self) -> None:
        if self._cur_p:
            # Avoid emitting whitespace-only paragraphs (common around lists/tables in QTextEdit HTML)
            txt = "".join((c.get("text", "") or "") for c in self._cur_p)
            # Keep intentional blank lines: QTextEdit represents them as <br>, which we map to "\n" chunks.
            has_explicit_break = any((c.get("text") == "\n") for c in self._cur_p)
            if txt.strip() != "" or has_explicit_break:
                self.items.append({"type": "p", "chunks": list(self._cur_p)})
        self._cur_p = []


def _html_to_docx_items(body_html: str) -> List[dict]:
    p = _HtmlBodyParser()
    p.feed(body_html or "")
    p.close()
    return p.items


def _set_table_borders(table) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    # Set borders for all cells in the table
    tbl = table._tbl
    for cell in tbl.iter_tceleds() if hasattr(tbl, 'iter_tceleds') else []:
        pass # Not used, we'll do it via tblPr
    
    # Simple way: add <w:tblBorders> to <w:tblPr>
    tblPr = tbl.xpath('w:tblPr')[0]
    borders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'</w:tblBorders>'
    )
    tblPr.append(borders)


def _to_roman(num: int) -> str:
    """Convert integer to Roman numeral"""
    roman_numerals = [
        (1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'),
        (100, 'C'), (90, 'XC'), (50, 'L'), (40, 'XL'),
        (10, 'X'), (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I')
    ]
    result = ""
    for value, numeral in roman_numerals:
        while num >= value:
            result += numeral
            num -= value
    return result


def _generate_signature_log(workflow: WorkflowService, doc: Document, current_user: User) -> str:
    """Generate encoded signature log for QR code"""
    from datetime import datetime
    
    # Helper function to encode employee info
    def encode_employee_info(employee_id: str, date_str: str) -> str:
        """Encode employee info as: (first half of employee_id) + day + year + month + (second half of employee_id)"""
        try:
            # Parse date string (assuming format like YYYY-MM-DD or similar)
            if '-' in date_str:
                parts = date_str.split('-')
                if len(parts) >= 3:
                    year = parts[0][-2:]  # Last 2 digits of year
                    month = parts[1].zfill(2)
                    day = parts[2].zfill(2)
                else:
                    # Fallback to current date
                    dt = datetime.now()
                    day = dt.strftime("%d")
                    year = dt.strftime("%y")
                    month = dt.strftime("%m")
            else:
                # Fallback to current date
                dt = datetime.now()
                day = dt.strftime("%d")
                year = dt.strftime("%y")
                month = dt.strftime("%m")
            
            emp_id_str = "".join(ch for ch in str(employee_id) if ch.isdigit())
            if not emp_id_str:
                emp_id_str = "0"

            # Split in the middle. For very short IDs (length 1) keep 1 digit in second_half.
            mid_point = max(1, len(emp_id_str) // 2)
            first_half = emp_id_str[:mid_point]
            second_half = emp_id_str[mid_point:]
            
            # Format: (first half of employee_id) + day + year + month + (second half of employee_id)
            return f"{first_half}{day}{year}{month}{second_half}"
        except Exception:
            # Fallback format if encoding fails - split the ID anyway
            emp_id_str = "".join(ch for ch in str(employee_id) if ch.isdigit())
            if not emp_id_str:
                emp_id_str = "0"
            mid_point = max(1, len(emp_id_str) // 2)
            first_half = emp_id_str[:mid_point]
            second_half = emp_id_str[mid_point:]
            return f"{first_half}010101{second_half}"
    
    # Start with initiator info
    encoded_parts = []
    
    # Get initiator info
    initiator_user = workflow.get_user(doc.created_by) if doc.created_by else current_user
    initiator_emp_id = (getattr(initiator_user, "employee_id", "") or "").strip()
    if not initiator_emp_id:
        initiator_emp_id = str(getattr(initiator_user, "id", doc.created_by or current_user.id))
    
    # Encode initiator with creation date
    creation_date = doc.created_at or datetime.now().isoformat()
    if 'T' in creation_date:
        creation_date = creation_date.split('T')[0]
    
    encoded_parts.append(encode_employee_info(initiator_emp_id, creation_date))
    
    # Add approvers if document has been sent for approval
    if doc.id:
        try:
            chain = workflow.get_approval_chain(doc.id)
            if chain:
                for step in chain:
                    approver_emp_id = ""
                    approver_user = workflow.get_user(step.user_id)
                    if approver_user is not None:
                        approver_emp_id = (getattr(approver_user, "employee_id", "") or "").strip()
                    if not approver_emp_id:
                        approver_emp_id = str(step.user_id)
                    # Use approval date if available, otherwise use current date
                    approval_date = step.approval_date if step.approval_date else datetime.now().isoformat()
                    if 'T' in approval_date:
                        approval_date = approval_date.split('T')[0]
                    
                    encoded_parts.append(encode_employee_info(approver_emp_id, approval_date))
        except Exception:
            pass
    
    # Join all encoded parts with a separator
    return "|".join(encoded_parts)


def _generate_qr_code(text: str) -> bytes:
    """Generate QR code image bytes from text"""
    import qrcode
    from io import BytesIO
    
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=2,
        border=1,
    )
    qr.add_data(text)
    qr.make(fit=True)
    
    img = qr.make_image(fill_color="black", back_color="white")
    
    # Resize to smaller size for document footer
    img = img.resize((80, 80))
    
    bio = BytesIO()
    img.save(bio, format='PNG')
    return bio.getvalue()


def _generate_docx_payload(
    template_path: Path,
    output_path: Path,
    workflow: WorkflowService,
    doc_model: Document,
    current_user: User,
    editor_html: str
) -> str:
    import docx
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = docx.Document(str(template_path))

    # Clear body paragraphs
    for paragraph in list(doc.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)

    # Set default font to Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Aggressively remove all spacing from default style
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.0
    paragraph_format.space_after_auto = False
    paragraph_format.space_before_auto = False
    
    # Ensure it applies correctly to all script types in Word via XML
    rPr = font.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')

    # 1. Ref and Date (Removed extra spacer before it)
    ref_value, date_str = _build_ref_and_date(workflow=workflow, doc=doc_model, user=current_user)
    ref_p = doc.add_paragraph()
    # Force Times New Roman on Ref line too
    r1 = ref_p.add_run(f"REF#: {ref_value}")
    r1.font.name = 'Times New Roman'
    r1._element.rPr.get_or_add_rFonts().set(qn('w:ascii'), 'Times New Roman')
    r1._element.rPr.get_or_add_rFonts().set(qn('w:hAnsi'), 'Times New Roman')

    ref_p.add_run("\t")
    
    r2 = ref_p.add_run(f"Date:{date_str}")
    r2.font.name = 'Times New Roman'
    r2._element.rPr.get_or_add_rFonts().set(qn('w:ascii'), 'Times New Roman')
    r2._element.rPr.get_or_add_rFonts().set(qn('w:hAnsi'), 'Times New Roman')

    try:
        section = doc.sections[0]
        # Move tab stop 0.8 inches left from the right margin to ensure it fits better in PDF
        tab_pos = section.page_width - section.right_margin - Inches(0.8)
        ref_p.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
    except Exception:
        ref_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)

    # 2. Subject
    subject_p = doc.add_paragraph()
    r_subj_label = subject_p.add_run("Subject: ")
    r_subj_label.bold = True
    r_subj_label.font.name = 'Times New Roman'
    
    r_subj_val = subject_p.add_run(doc_model.subject or "")
    r_subj_val.font.name = 'Times New Roman'

    # 3. Content
    try:
        username_set = {
            (u.username or "").strip().lower() for u in (workflow.list_users() or []) if getattr(u, "username", None)
        }
    except Exception:
        username_set = set()

    items = _html_to_docx_items(editor_html)
    
    def apply_chunk_to_paragraph(p, chunk):
        text = chunk.get("text", "")
        if not text: return
        # Filter unwanted system names etc
        for line in _filtered_lines_from_text(text, username_set):
            r = p.add_run(line)
            r.bold = chunk.get("bold", False)
            r.italic = chunk.get("italic", False)
            r.underline = chunk.get("underline", False)
            f_name = chunk.get("font", "Times New Roman")
            r.font.name = f_name
            r.font.size = Pt(chunk.get("size", 12))
            
            # Force font via XML to avoid DengXian/Theme overrides
            rFonts = r._element.get_or_add_rPr().get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), f_name)
            rFonts.set(qn('w:hAnsi'), f_name)
            rFonts.set(qn('w:eastAsia'), f_name)
            rFonts.set(qn('w:cs'), f_name)

    for it in items:
        if it.get("type") == "p":
            p = doc.add_paragraph()
            # Apply aggressive spacing control to match editor
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after_auto = False
            p.paragraph_format.space_before_auto = False
            # Force spacing via XML to override template defaults
            pPr = p._element.get_or_add_pPr()
            pPr.set(qn('w:spaceBefore'), '0')
            pPr.set(qn('w:spaceAfter'), '0')
            pPr.set(qn('w:line'), '240')  # Single line spacing
            pPr.set(qn('w:lineRule'), 'auto')
            for chunk in it.get("chunks", []):
                apply_chunk_to_paragraph(p, chunk)
        elif it.get("type") == "list":
            list_type = it.get("list_type", "bullet")
            list_items = it.get("items", [])
            indent_pt = float(it.get("indent_pt", 0.0) or 0.0)
            
            # Initialize counters for different list types
            if not hasattr(_generate_docx_payload, '_list_counters'):
                _generate_docx_payload._list_counters = {}
            list_counter_key = f"{id(doc)}_{list_type}"
            # Restart numbering per list block
            _generate_docx_payload._list_counters[list_counter_key] = 1
            
            for item in list_items:
                p = doc.add_paragraph()
                # Apply aggressive spacing control for list items to match editor
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after_auto = False
                p.paragraph_format.space_before_auto = False
                if indent_pt > 0:
                    p.paragraph_format.left_indent = Pt(indent_pt)
                # Force spacing via XML to override template defaults
                pPr = p._element.get_or_add_pPr()
                pPr.set(qn('w:spaceBefore'), '0')
                pPr.set(qn('w:spaceAfter'), '0')
                pPr.set(qn('w:line'), '240')  # Single line spacing
                pPr.set(qn('w:lineRule'), 'auto')
                # Manually create list formatting based on type
                if list_type == "bullet":
                    run = p.add_run("• ")
                elif list_type == "alpha":
                    # Convert number to letter (a, b, c, ...)
                    letter = chr(ord('a') + _generate_docx_payload._list_counters[list_counter_key] - 1)
                    run = p.add_run(f"{letter}. ")
                elif list_type == "roman":
                    # Convert number to Roman numeral (I, II, III, ...)
                    num = _generate_docx_payload._list_counters[list_counter_key]
                    roman = _to_roman(num)
                    run = p.add_run(f"{roman}. ")
                else:  # number
                    run = p.add_run(f"{_generate_docx_payload._list_counters[list_counter_key]}. ")
                
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                _generate_docx_payload._list_counters[list_counter_key] += 1
                
                for chunk in item.get("chunks", []):
                    apply_chunk_to_paragraph(p, chunk)
        elif it.get("type") == "table":
            rows_data = it.get("rows") or []
            if not rows_data: continue
            num_rows = len(rows_data)
            num_cols = max(len(r) for r in rows_data)
            table = doc.add_table(rows=num_rows, cols=num_cols)
            try:
                table.style = 'Table Grid'
            except Exception:
                _set_table_borders(table)
            
            for r_idx, row_data in enumerate(rows_data):
                for c_idx, cell_data_list in enumerate(row_data):
                    if c_idx < num_cols:
                        cell_p = table.cell(r_idx, c_idx).paragraphs[0]
                        for chunk in cell_data_list:
                            apply_chunk_to_paragraph(cell_p, chunk)


    # 4. Initiator
    initiator_user = workflow.get_user(doc_model.created_by) if doc_model.created_by else current_user
    initiator_name = getattr(initiator_user, "full_name", "").strip() or getattr(initiator_user, "name", "").strip()
    initiator_designation = getattr(initiator_user, "designation", "").strip()
    if initiator_name:
        doc.add_paragraph()
        sig = getattr(doc_model, "initiator_signature_png", None)
        if sig:
            try:
                sig_p = doc.add_paragraph()
                sig_p.paragraph_format.alignment = WD_TAB_ALIGNMENT.RIGHT
                run = sig_p.add_run()
                run.add_picture(io.BytesIO(sig), height=Inches(0.5))
            except Exception:
                pass
        # Combine name and designation in separate lines with minimal spacing
        name_content = initiator_name
        if initiator_designation:
            name_content += f"\n{initiator_designation}"
        
        name_p = doc.add_paragraph(name_content)
        name_p.paragraph_format.alignment = WD_TAB_ALIGNMENT.RIGHT
        # Reduce spacing between name and designation lines
        name_p.paragraph_format.space_after = Pt(0)
        name_p.paragraph_format.space_before = Pt(0)
        name_p.paragraph_format.line_spacing = 1.0

    # 5. Approval Chain
    try:
        chain = workflow.get_approval_chain(doc_model.id) if doc_model.id else []
        if chain:
            doc.add_paragraph()
            for step in chain:
                u = workflow.get_user(int(step.user_id))
                if u is None: continue
                desig = (getattr(u, "designation", "") or "").strip()
                if desig:
                    p = doc.add_paragraph()
                    
                    # 1. Designation First (Bold)
                    r_desig = p.add_run(desig)
                    r_desig.bold = True
                    r_desig.font.name = 'Times New Roman'
                    r_desig.font.size = Pt(12)
                    # Force font
                    rFonts = r_desig._element.get_or_add_rPr().get_or_add_rFonts()
                    rFonts.set(qn('w:ascii'), 'Times New Roman')
                    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    
                    p.add_run("  ")

                    # 2. Signature Second (Shifted down)
                    p.add_run("   ")
                    sig = getattr(step, "signature_png", None)
                    if sig:
                        try:
                            run = p.add_run()
                            run.add_picture(io.BytesIO(sig), height=Inches(0.35))
                            
                            # Shift down via XML w:position (negative value is down, units are half-points)
                            # 1 pt = 2 half-points. Let's shift by ~12 half-points (~6pt)
                            rPr = run._element.get_or_add_rPr()
                            pos = rPr.get_or_add_position()
                            pos.set(qn('w:val'), '-12')
                        except Exception: pass
    except Exception: pass

    # 6. Add QR Code to the footer
    try:
        signature_log_text = _generate_signature_log(workflow, doc_model, current_user)
        qr_bytes = _generate_qr_code(signature_log_text)
        
        # Get the first section
        section = doc.sections[0]
        footer = section.footer
        
        # Clear any existing footer content
        for paragraph in footer.paragraphs:
            paragraph.clear()
        
        # Add QR code to footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_TAB_ALIGNMENT.CENTER
        
        # Add QR code image to footer
        qr_run = footer_para.add_run()
        qr_run.add_picture(io.BytesIO(qr_bytes), width=Inches(0.8))
        
        # Set footer font properties
        footer_para.runs[0].font.size = Pt(8)
        
    except Exception:
        # Silently fail if QR code generation fails
        pass

    doc.save(str(output_path))
    return str(output_path)


def _convert_docx_to_pdf_win32(docx_path: str) -> str:
    if not win32com:
        raise ImportError("pywin32 (win32com) is not available for PDF conversion.")
    
    # Try to initialize COM for the current thread
    try:
        import pythoncom
        pythoncom.CoInitialize()
    except Exception:
        pass

    pdf_path = docx_path.rsplit(".", 1)[0] + ".pdf"
    abs_docx = str(Path(docx_path).resolve())
    abs_pdf = str(Path(pdf_path).resolve())
    
    word = None
    try:
        # Try a fresh instance
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0 
        
        doc_obj = word.Documents.Open(abs_docx, ReadOnly=True, ConfirmConversions=False)
        
        try:
            # 1. Try ExportAsFixedFormat (Highest fidelity)
            # wdExportFormatPDF = 17; wdExportOptimizeForPrint = 0
            doc_obj.ExportAsFixedFormat(
                OutputFileName=abs_pdf,
                ExportFormat=17,
                OpenAfterExport=False,
                OptimizeFor=0, 
                Range=0, 
                Item=0,
                IncludeDocProps=True,
                KeepIRM=True,
                CreateBookmarks=0,
                DocStructureTags=True,
                BitmapMissingFonts=True,
                UseISO19005_1=False
            )
        except Exception:
            # 2. Fallback to SaveAs (Compatible with older Word and WPS Office)
            # wdFormatPDF = 17
            doc_obj.SaveAs(abs_pdf, FileFormat=17)
            
        doc_obj.Close(0) # wdDoNotSaveChanges
    except Exception as e:
        # Final error wrap
        msg = str(e) if str(e).strip() else "Unknown COM error"
        raise RuntimeError(f"Word/WPS PDF conversion failed: {msg}")
    finally:
        if word:
            try:
                word.Quit()
            except Exception:
                pass
    return pdf_path


def _filtered_lines_from_text(text: str, username_set: set[str]) -> List[str]:
    out: List[str] = []
    for ln in (text or "").splitlines():
        if _is_signature_artifact_line(ln):
            continue
        if username_set and ln.strip().lower() in username_set:
            continue
        out.append(ln)
    return out


class AdvancedTextEdit(QTextEdit):
    def __init__(self, parent=None):
        super().__init__(parent)

    def _adjust_current_list_indent(self, delta: int) -> bool:
        cursor = self.textCursor()
        lst = cursor.currentList()
        if not lst:
            return False

        fmt = lst.format()
        new_indent = max(0, int(fmt.indent()) + int(delta))
        fmt.setIndent(new_indent)
        lst.setFormat(fmt)
        return True

    def keyPressEvent(self, event):
        key = event.key()
        if key == Qt.Key_Tab:
            if self._adjust_current_list_indent(1):
                return
        if key == Qt.Key_Backtab:
            if self._adjust_current_list_indent(-1):
                return
        super().keyPressEvent(event)

    def _get_editor_window(self):
        p = self.parent()
        while p:
            if hasattr(p, "_add_row_above"):
                return p
            p = p.parent()
        return None

    def contextMenuEvent(self, event):
        try:
            menu = self.createStandardContextMenu()
            cursor = self.textCursor()
            table = cursor.currentTable()
            win = self._get_editor_window()

            if table and win:
                menu.addSeparator()
                table_menu = menu.addMenu("Table Tools")

                insert_row_above = table_menu.addAction("Insert Row Above")
                insert_row_below = table_menu.addAction("Insert Row Below")
                insert_row_above.triggered.connect(win._add_row_above)
                insert_row_below.triggered.connect(win._add_row_below)

                table_menu.addSeparator()
                insert_col_before = table_menu.addAction("Insert Column Before")
                insert_col_after = table_menu.addAction("Insert Column After")
                insert_col_before.triggered.connect(win._add_column_before)
                insert_col_after.triggered.connect(win._add_column_after)

                table_menu.addSeparator()
                delete_row = table_menu.addAction("Delete Row")
                delete_col = table_menu.addAction("Delete Column")
                delete_row.triggered.connect(win._remove_row)
                delete_col.triggered.connect(win._remove_column)

            menu.exec_(event.globalPos())
        except RuntimeError:
            # Object deleted already, just ignore the event
            pass


class _SignatureCanvas(QWidget):
    def __init__(self, parent=None) -> None:
        super().__init__(parent)
        self._img = QImage(520, 180, QImage.Format_ARGB32)
        self._img.fill(Qt.white)
        self._last: Optional[QPoint] = None
        self.setMinimumSize(520, 180)

    def clear(self) -> None:
        self._img.fill(Qt.white)
        self.update()

    def to_png_bytes(self) -> bytes:
        ba = QByteArray()
        buf = QBuffer(ba)
        buf.open(QIODevice.WriteOnly)
        self._img.save(buf, "PNG")
        return bytes(ba.data())

    def paintEvent(self, event) -> None:
        p = QPainter(self)
        p.drawImage(0, 0, self._img)

    def mousePressEvent(self, event) -> None:
        if event.button() == Qt.LeftButton:
            self._last = event.pos()

    def mouseMoveEvent(self, event) -> None:
        if self._last is None:
            return
        p = QPainter(self._img)
        pen = QPen(QColor(0, 0, 0))
        pen.setWidth(3)
        pen.setCapStyle(Qt.RoundCap)
        p.setPen(pen)
        p.drawLine(self._last, event.pos())
        self._last = event.pos()
        self.update()

    def mouseReleaseEvent(self, event) -> None:
        self._last = None


class _SignatureDialog(QDialog):
    def __init__(self, parent=None) -> None:
        super().__init__(parent)
        self.setWindowTitle("Signature")
        self.setModal(True)
        self.resize(600, 360)

        self._mode_draw = QPushButton("Draw with Mouse")
        self._mode_upload = QPushButton("Upload Image")

        self._canvas = _SignatureCanvas(self)
        self._upload_label = QLabel("No image selected")
        self._upload_label.setWordWrap(True)
        self._uploaded_png: Optional[bytes] = None

        self._browse_btn = QPushButton("Browse...")
        self._browse_btn.clicked.connect(self._on_browse)

        self._clear_btn = QPushButton("Clear")
        self._clear_btn.clicked.connect(self._canvas.clear)

        self._mode = "draw"
        self._mode_draw.setEnabled(False)
        self._mode_upload.clicked.connect(self._set_mode_upload)
        self._mode_draw.clicked.connect(self._set_mode_draw)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)

        top = QHBoxLayout()
        top.addWidget(self._mode_draw)
        top.addWidget(self._mode_upload)
        top.addStretch(1)

        upload_row = QHBoxLayout()
        upload_row.addWidget(self._browse_btn)
        upload_row.addWidget(self._upload_label, 1)

        layout = QVBoxLayout()
        layout.addLayout(top)
        layout.addWidget(self._canvas)
        layout.addWidget(self._clear_btn)
        layout.addLayout(upload_row)
        layout.addWidget(buttons)
        self.setLayout(layout)

        self._apply_mode_visibility()

    def _set_mode_draw(self) -> None:
        self._mode = "draw"
        self._mode_draw.setEnabled(False)
        self._mode_upload.setEnabled(True)
        self._apply_mode_visibility()

    def _set_mode_upload(self) -> None:
        self._mode = "upload"
        self._mode_upload.setEnabled(False)
        self._mode_draw.setEnabled(True)
        self._apply_mode_visibility()

    def _apply_mode_visibility(self) -> None:
        is_draw = self._mode == "draw"
        self._canvas.setVisible(is_draw)
        self._clear_btn.setVisible(is_draw)
        self._browse_btn.setVisible(not is_draw)
        self._upload_label.setVisible(not is_draw)

    def _on_browse(self) -> None:
        path, _ = QFileDialog.getOpenFileName(self, "Select signature image", str(Path.home()), "Images (*.png *.jpg *.jpeg *.bmp)")
        if not path:
            return
        img = QImage(path)
        if img.isNull():
            QMessageBox.warning(self, "Signature", "Failed to load image")
            return
        # Normalize to PNG bytes
        ba = QByteArray()
        buf = QBuffer(ba)
        buf.open(QIODevice.WriteOnly)
        img.save(buf, "PNG")
        self._uploaded_png = bytes(ba.data())
        self._upload_label.setText(path)

    def signature_png_bytes(self) -> Optional[bytes]:
        if self._mode == "upload":
            return self._uploaded_png
        # draw
        return self._canvas.to_png_bytes()


class _ChainDialog(QDialog):
    def __init__(self, workflow: WorkflowService, current_user: User, initial_user_ids: List[int], parent=None) -> None:
        super().__init__(parent)
        self._workflow = workflow
        self._current_user = current_user

        self.setWindowTitle("Approval Chain")
        self.setModal(True)
        self.resize(520, 340)

        self._users_combo = QComboBox(self)
        for u in self._workflow.list_users():
            self._users_combo.addItem(u.display_label(), u.id)

        from PyQt5.QtWidgets import QListWidget

        self._list = QListWidget(self)
        for uid in initial_user_ids:
            self._append_user(uid)

        add_btn = QPushButton("Add")
        remove_btn = QPushButton("Remove")
        up_btn = QPushButton("Up")
        down_btn = QPushButton("Down")

        add_btn.clicked.connect(self._add_selected)
        remove_btn.clicked.connect(self._remove_selected)
        up_btn.clicked.connect(lambda: self._move(-1))
        down_btn.clicked.connect(lambda: self._move(1))

        controls = QHBoxLayout()
        controls.addWidget(QLabel("User:"))
        controls.addWidget(self._users_combo)
        controls.addWidget(add_btn)
        controls.addStretch(1)

        reorder = QHBoxLayout()
        reorder.addWidget(remove_btn)
        reorder.addWidget(up_btn)
        reorder.addWidget(down_btn)
        reorder.addStretch(1)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)

        layout = QVBoxLayout()
        layout.addLayout(controls)
        layout.addWidget(self._list)
        layout.addLayout(reorder)
        layout.addWidget(buttons)
        self.setLayout(layout)

    def _append_user(self, user_id: int) -> None:
        u = self._workflow.get_user(user_id)
        label = str(user_id) if u is None else u.display_label()
        item = QListWidgetItem(label)
        item.setData(Qt.UserRole, int(user_id))
        self._list.addItem(item)

    def _add_selected(self) -> None:
        uid = self._users_combo.currentData()
        if uid is None:
            return
        self._append_user(int(uid))

    def _remove_selected(self) -> None:
        row = self._list.currentRow()
        if row >= 0:
            self._list.takeItem(row)

    def _move(self, delta: int) -> None:
        row = self._list.currentRow()
        if row < 0:
            return
        new_row = row + delta
        if new_row < 0 or new_row >= self._list.count():
            return
        item = self._list.takeItem(row)
        self._list.insertItem(new_row, item)
        self._list.setCurrentRow(new_row)

    def user_ids(self) -> List[int]:
        out: List[int] = []
        for i in range(self._list.count()):
            item = self._list.item(i)
            out.append(int(item.data(Qt.UserRole)))
        return out


class _LetterPreviewDialog(QDialog):
    def __init__(self, workflow: WorkflowService, current_user: User, doc: Document, parent=None) -> None:
        super().__init__(parent)
        self._workflow = workflow
        self._current_user = current_user
        self._doc = doc
        self.setWindowTitle("Letter Preview")
        self.setModal(True)
        self.resize(800, 900)

        # Build HTML for preview (header + ref/date + subject + body + initiator)
        header_path = Path(__file__).resolve().parents[1] / "documents" / "header.png"
        header_html = ""
        if header_path.exists():
            header_html = f"<div style='text-align:center; margin-bottom:24px;'><img src='{header_path.as_uri()}' style='max-width:600px; height:auto;'/></div>"

        # REF#/Date line (Adjusted to be inward slightly from edges)
        ref_value, date_str = _build_ref_and_date(workflow=self._workflow, doc=self._doc, user=self._current_user)
        ref_html = f"<div style='margin:0 auto; max-width:540px; margin-bottom:12px; display:flex; justify-content:space-between;'><span>REF#: {ref_value}</span><span>Date:{date_str}</span></div>"

        subject_html = f"<div style='margin:0 auto; max-width:600px; margin-bottom:12px;'><b>Subject:</b> {html.escape(doc.subject or '')}</div>"

        # Render the editor HTML directly so tables/images/formatting are preserved.
        body_html = f"<div style='margin:0 auto; max-width:600px; white-space:normal;'>{doc.content or ''}</div>"

        # Initiator name and designation at the end, right-aligned (use original creator)
        initiator_user = self._workflow.get_user(self._doc.created_by) if self._doc.created_by else self._current_user
        initiator_name = getattr(initiator_user, "full_name", "").strip() or getattr(initiator_user, "name", "").strip()
        initiator_designation = getattr(initiator_user, "designation", "").strip()
        initiator_html = ""
        if initiator_name:
            initiator_html = "<div style='margin:0 auto; max-width:600px; margin-top:24px; text-align:right;'>"
            sig = getattr(self._doc, "initiator_signature_png", None)
            if sig:
                b64 = base64.b64encode(sig).decode("ascii")
                # Use a block wrapper + explicit break to force the signature above the name in QTextEdit.
                initiator_html += (
                    f"<div style='display:block;'><img src='data:image/png;base64,{b64}' "
                    "style='height:60px; width:auto; display:block; margin-left:auto; margin-bottom:6px;'/></div>"
                )
            initiator_html += f"<div style='display:block;'>{initiator_name}</div>"
            if initiator_designation:
                initiator_html += f"<div style='display:block;'>{initiator_designation}</div>"
            initiator_html += "</div>"

        # Approval chain designations + signatures, left-aligned with line spacing
        approvers_html = ""
        try:
            chain = self._workflow.get_approval_chain(self._doc.id) if self._doc.id else []
            rows = []
            for step in chain:
                u = self._workflow.get_user(int(step.user_id))
                if u is None:
                    continue
                desig = (getattr(u, "designation", "") or "").strip()
                if desig:
                    img_html = ""
                    sig = getattr(step, "signature_png", None)
                    if sig:
                        b64 = base64.b64encode(sig).decode("ascii")
                        img_html = f"<img src='data:image/png;base64,{b64}' style='height:40px; width:auto; vertical-align: -12px; margin-left: 8px;'/>"
                    row_html = (
                        "<div style='display:flex; align-items:baseline; margin-bottom:4px;'>"
                        f"<span style='font-weight:bold;'>{html.escape(desig)}</span>{img_html}"
                        "</div>"
                    )
                    rows.append(row_html)

            if rows:
                approvers_html = "<div style='margin:0 auto; max-width:600px; margin-top:24px;'>" + "".join(rows) + "</div>"
        except Exception:
            pass

        full_html = f"<html><body style='font-family:Arial; padding:20px;'>{header_html}{ref_html}{subject_html}{body_html}{initiator_html}{approvers_html}</body></html>"

        self._preview = QTextEdit(self)
        self._preview.setAcceptRichText(True)
        self._preview.setHtml(full_html)
        self._preview.setReadOnly(True)

        self._save_pdf_btn = QPushButton("Save letter", self)
        self._save_pdf_btn.clicked.connect(self._save_pdf)

        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        btn_layout.addWidget(self._save_pdf_btn)

        layout = QVBoxLayout()
        layout.addWidget(self._preview)
        layout.addLayout(btn_layout)
        self.setLayout(layout)

    def _save_pdf(self) -> None:
        out_dir = Path.home() / "Documents"
        out_dir.mkdir(parents=True, exist_ok=True)
        safe_title = (self._doc.title or "Untitled").strip().replace("\\", "_").replace("/", "_")
        docx_path = out_dir / f"{safe_title}.docx"
        
        # Use the admin-uploaded template in letterhead folder if it exists
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        template_path = os.path.join(base_dir, "letterhead", "REF.docx")
        if not os.path.exists(template_path):
            # Fallback to default
            template_path = os.path.join(base_dir, "documents", "REF.docx")
        
        try:
            # 1. Generate DOCX
            _generate_docx_payload(
                template_path=Path(template_path),
                output_path=docx_path,
                workflow=self._workflow,
                doc_model=self._doc,
                current_user=self._current_user,
                editor_html=self._doc.content
            )
            
            # Ask if user wants PDF too
            res = QMessageBox.question(
                self, "Export", 
                f"DOCX saved to {docx_path}\n\nWould you like to export a PDF version as well?",
                QMessageBox.Yes | QMessageBox.No
            )
            
            if res == QMessageBox.Yes:
                pdf_path = _convert_docx_to_pdf_win32(str(docx_path))
                QMessageBox.information(self, "Success", f"PDF exported to:\n{pdf_path}")
            else:
                QMessageBox.information(self, "Success", f"DOCX exported to:\n{docx_path}")
                
        except Exception as e:
            traceback.print_exc()
            QMessageBox.critical(self, "Export failed", f"Failed to export: {e}")


class EditorWindow(QMainWindow):
    saved = pyqtSignal()

    def __init__(
        self,
        workflow: WorkflowService,
        comment_service: CommentService,
        current_user: User,
        document: Optional[Document],
        parent: QWidget = None
    ) -> None:
        super().__init__(parent)
        
        # Make the window modal to block the main dashboard
        self.setWindowModality(Qt.WindowModal)

        self._workflow = workflow
        self._comments = comment_service
        self._current_user = current_user

        self._doc: Document = document or Document(
            id=None,
            title="Untitled",
            subject="",
            content=_load_ref_body_html(),
            created_by=current_user.id,
            status="Draft",
            assigned_to=None,
            current_step=0,
        )
        self._dirty = False

        self._pending_chain_user_ids: List[int] = []

        self.setWindowTitle(self._window_title())
        self.resize(1000, 700)
        self.setWindowState(Qt.WindowMaximized)

        self._header_label = QLabel(self)
        self._header_label.setAlignment(Qt.AlignHCenter)
        self._header_label.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        self._header_label.setMinimumHeight(60)
        self._header_label.setMaximumHeight(150)
        header_path = Path(__file__).resolve().parents[1] / "documents" / "header.png"
        if header_path.exists():
            pm = QPixmap(str(header_path))
            self._header_pixmap = pm
            self._header_label.setPixmap(pm)
        else:
            self._header_pixmap = QPixmap()
            self._header_label.setText("")

        self._title = QLineEdit(self)
        self._title.setText(self._doc.title)

        self._subject = QLineEdit(self)
        self._subject.setText(getattr(self._doc, 'subject', ''))
        self._subject.setPlaceholderText("Enter subject")

        self._editor = AdvancedTextEdit(self)
        self._editor.setAcceptRichText(True)
        self._editor.setHtml(self._doc.content or "")
        self._editor.setObjectName("qt_editor")

        self._title.textChanged.connect(self._on_changed)
        self._subject.textChanged.connect(self._on_changed)
        self._editor.textChanged.connect(self._on_changed)
        self._editor.cursorPositionChanged.connect(self._update_format_controls)

        self._create_menus()
        self._create_toolbar()

        self._save_btn = QPushButton("Save")
        self._view_btn = QPushButton("View Letter")
        self._delete_btn = QPushButton("Delete")
        self._send_btn = QPushButton("Send for Approval")
        self._approve_btn = QPushButton("Approve")
        self._reject_btn = QPushButton("Reject")
        self._export_pdf_btn = QPushButton("Export PDF")

        self._save_btn.clicked.connect(self.save)
        self._view_btn.clicked.connect(self._view_letter)
        self._delete_btn.clicked.connect(self._delete_document)
        self._send_btn.clicked.connect(self.send_for_approval)
        self._approve_btn.clicked.connect(self.approve)
        self._reject_btn.clicked.connect(self.reject)
        self._export_pdf_btn.clicked.connect(self.export_pdf)

        btn_bar = QHBoxLayout()
        btn_bar.addWidget(self._save_btn)
        btn_bar.addWidget(self._view_btn)
        btn_bar.addWidget(self._export_pdf_btn)
        btn_bar.addWidget(self._delete_btn)
        btn_bar.addWidget(self._send_btn)
        btn_bar.addStretch(1)
        btn_bar.addWidget(self._approve_btn)
        btn_bar.addWidget(self._reject_btn)

        self._chain_label = QLabel("Approval: (not configured)")
        self._chain_label.setObjectName("SectionTitle")
        self._chain_label.hide()

        self._comment_widget = CommentWidget(self._on_comment_added, self)
        self._comment_widget.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Expanding)

        form = QFormLayout()
        form.addRow("Title:", self._title)
        form.addRow("Subject:", self._subject)

        body_label = QLabel("Body")
        body_label.setObjectName("SectionTitle")

        approval_label = QLabel("Approval Chain")
        approval_label.setObjectName("SectionTitle")

        self._approval_chain_btn = QPushButton("Set Approvers")
        self._approval_chain_btn.clicked.connect(self.configure_chain)

        self._approval_chain_value = QLabel("")
        self._approval_chain_value.setWordWrap(True)

        approval_row = QHBoxLayout()
        approval_row.addWidget(self._approval_chain_btn)
        approval_row.addStretch(1)

        left = QWidget(self)
        left_layout = QVBoxLayout()
        left_layout.addLayout(form)
        left_layout.addWidget(self._header_label)
        left_layout.addWidget(body_label)
        left_layout.addWidget(self._editor)
        left_layout.addWidget(approval_label)
        left_layout.addLayout(approval_row)

        left_layout.addWidget(self._approval_chain_value)
        left_layout.addLayout(btn_bar)
        left.setLayout(left_layout)

        right = QWidget(self)
        right.setFixedWidth(340)
        right_layout = QVBoxLayout()
        right_layout.addWidget(self._view_btn)
        right_layout.addWidget(self._comment_widget, 1)
        right.setLayout(right_layout)

        root = QWidget(self)
        main_layout = QHBoxLayout()
        main_layout.addWidget(left, 1)
        main_layout.addWidget(right)
        root.setLayout(main_layout)
        self.setCentralWidget(root)

        self.statusBar()
        self._sync_action_visibility()
        self._update_status()
        self._refresh_approval_chain_section()
        self._reload_comments()

    def _create_menus(self) -> None:
        mb = self.menuBar()

        file_menu = mb.addMenu("File")
        act_save = QAction("Save", self)
        act_save.setShortcut("Ctrl+S")
        act_save.triggered.connect(self.save)
        file_menu.addAction(act_save)

        act_export_pdf = QAction("Export PDF", self)
        act_export_pdf.triggered.connect(self.export_pdf)
        file_menu.addAction(act_export_pdf)

        act_view = QAction("View Letter", self)
        act_view.triggered.connect(self._view_letter)
        file_menu.addAction(act_view)

        file_menu.addSeparator()
        act_close = QAction("Close", self)
        act_close.setShortcut("Ctrl+W")
        act_close.triggered.connect(self.close)
        file_menu.addAction(act_close)

        edit_menu = mb.addMenu("Edit")
        act_undo = QAction("Undo", self)
        act_undo.setShortcut("Ctrl+Z")
        act_undo.triggered.connect(self._editor.undo)
        edit_menu.addAction(act_undo)

        act_redo = QAction("Redo", self)
        act_redo.setShortcut("Ctrl+Y")
        act_redo.triggered.connect(self._editor.redo)
        edit_menu.addAction(act_redo)

        edit_menu.addSeparator()
        act_cut = QAction("Cut", self)
        act_cut.setShortcut("Ctrl+X")
        act_cut.triggered.connect(self._editor.cut)
        edit_menu.addAction(act_cut)

        act_copy = QAction("Copy", self)
        act_copy.setShortcut("Ctrl+C")
        act_copy.triggered.connect(self._editor.copy)
        edit_menu.addAction(act_copy)

        act_paste = QAction("Paste", self)
        act_paste.setShortcut("Ctrl+V")
        act_paste.triggered.connect(self._editor.paste)
        edit_menu.addAction(act_paste)

        insert_menu = mb.addMenu("Insert")
        act_insert_table = QAction("Table", self)
        act_insert_table.triggered.connect(self._insert_table)
        insert_menu.addAction(act_insert_table)

        act_table_size = QAction("Table Size", self)
        act_table_size.triggered.connect(self._resize_table)
        insert_menu.addAction(act_table_size)

        format_menu = mb.addMenu("Format")
        act_bold = QAction("Bold", self)
        act_bold.setShortcut("Ctrl+B")
        act_bold.triggered.connect(lambda: self._toggle_bold(True))
        format_menu.addAction(act_bold)

        act_italic = QAction("Italic", self)
        act_italic.setShortcut("Ctrl+I")
        act_italic.triggered.connect(lambda: self._toggle_italic(True))
        format_menu.addAction(act_italic)

        act_underline = QAction("Underline", self)
        act_underline.setShortcut("Ctrl+U")
        act_underline.triggered.connect(lambda: self._toggle_underline(True))
        format_menu.addAction(act_underline)

        format_menu.addSeparator()

        list_menu = format_menu.addMenu("List")
        act_list_none = QAction("None", self)
        act_list_none.triggered.connect(lambda: self._list_combo.setCurrentIndex(0))
        list_menu.addAction(act_list_none)

        act_list_bullet = QAction("Bullets", self)
        act_list_bullet.triggered.connect(lambda: self._list_combo.setCurrentIndex(1))
        list_menu.addAction(act_list_bullet)

        act_list_number = QAction("Numbers (1, 2, 3)", self)
        act_list_number.triggered.connect(lambda: self._list_combo.setCurrentIndex(2))
        list_menu.addAction(act_list_number)

        act_list_alpha = QAction("Letters (a, b, c)", self)
        act_list_alpha.triggered.connect(lambda: self._list_combo.setCurrentIndex(3))
        list_menu.addAction(act_list_alpha)

        act_list_roman = QAction("Roman (I, II, III)", self)
        act_list_roman.triggered.connect(lambda: self._list_combo.setCurrentIndex(4))
        list_menu.addAction(act_list_roman)

        act_indent_inc = QAction("Increase Indent", self)
        act_indent_inc.setShortcut("Ctrl+]")
        act_indent_inc.triggered.connect(lambda: self._editor._adjust_current_list_indent(1))
        format_menu.addAction(act_indent_inc)

        act_indent_dec = QAction("Decrease Indent", self)
        act_indent_dec.setShortcut("Ctrl+[")
        act_indent_dec.triggered.connect(lambda: self._editor._adjust_current_list_indent(-1))
        format_menu.addAction(act_indent_dec)

        workflow_menu = mb.addMenu("Workflow")
        act_set_approvers = QAction("Set Approvers", self)
        act_set_approvers.triggered.connect(self.configure_chain)
        workflow_menu.addAction(act_set_approvers)

        workflow_menu.addSeparator()
        act_send = QAction("Send for Approval", self)
        act_send.triggered.connect(self.send_for_approval)
        workflow_menu.addAction(act_send)

        act_approve = QAction("Approve", self)
        act_approve.triggered.connect(self.approve)
        workflow_menu.addAction(act_approve)

        act_reject = QAction("Reject", self)
        act_reject.triggered.connect(self.reject)
        workflow_menu.addAction(act_reject)

        profile_menu = mb.addMenu("Profile")
        act_change_sig = QAction("Change Signature", self)
        act_change_sig.triggered.connect(self._change_signature)
        profile_menu.addAction(act_change_sig)

    def _change_signature(self) -> None:
        dlg = _SignatureDialog(self)
        if dlg.exec_() != QDialog.Accepted:
            return
        sig_png = dlg.signature_png_bytes()
        if not sig_png:
            return
        try:
            self._workflow.set_user_signature_png(user_id=self._current_user.id, signature_png=sig_png)
        except Exception as exc:
            QMessageBox.critical(self, "Failed", str(exc))

    def resizeEvent(self, event) -> None:
        super().resizeEvent(event)
        if self._header_pixmap.isNull():
            return
        # Limit max width to keep image crisp and professional
        max_w = 600
        w = min(max_w, max(1, self._header_label.width()))
        scaled = self._header_pixmap.scaledToWidth(w, Qt.SmoothTransformation)
        self._header_label.setPixmap(scaled)

    def _export_to_docx_template(self) -> str:
        out_dir = Path.home() / "Documents"
        out_dir.mkdir(parents=True, exist_ok=True)
        safe_title = (self._doc.title or "Untitled").strip().replace("\\", "_").replace("/", "_")
        out_path = out_dir / f"{safe_title}.docx"
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        template_path = os.path.join(base_dir, "letterhead", "REF.docx")
        if not os.path.exists(template_path):
            template_path = os.path.join(base_dir, "documents", "REF.docx")

        return _generate_docx_payload(
            template_path=Path(template_path),
            output_path=out_path,
            workflow=self._workflow,
            doc_model=self._doc,
            current_user=self._current_user,
            editor_html=self._editor.toHtml()
        )

    def _convert_docx_to_pdf(self, docx_path: str) -> str:
        return _convert_docx_to_pdf_win32(docx_path)

    def export_pdf(self) -> None:
        self._sync_doc_from_ui()
        if self._doc.id is None or self._dirty:
            QMessageBox.warning(self, "Save required", "Please save the document before exporting to PDF.")
            return

        try:
            # 1. Generate DOCX from template
            docx_path = self._export_to_docx_template()
            
            # 2. Convert to PDF via Word
            pdf_path = self._convert_docx_to_pdf(docx_path)
            
            QMessageBox.information(self, "Success", f"Document exported as:\nDOCX: {docx_path}\nPDF: {pdf_path}")
        except Exception as e:
            traceback.print_exc()
            QMessageBox.critical(self, "Export failed", f"Failed to export PDF:\n{e}")

    def _view_letter(self) -> None:
        # Sync current UI values into the document model for preview/export
        self._sync_doc_from_ui()
        dlg = _LetterPreviewDialog(self._workflow, self._current_user, self._doc, self)
        dlg.exec_()

    def _delete_document(self) -> None:
        if self._doc.id is None:
            QMessageBox.information(self, "Info", "This document has not been saved yet.")
            return
        reply = QMessageBox.question(
            self,
            "Delete Document",
            f"Are you sure you want to delete '{self._doc.title}'?\nThis action cannot be undone.",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No,
        )
        if reply != QMessageBox.Yes:
            return
        try:
            self._workflow.delete_document(self._doc.id)
            QMessageBox.information(self, "Deleted", "Document deleted successfully.")
            self.saved.emit()
            self.close()
        except Exception as exc:
            QMessageBox.critical(self, "Delete failed", str(exc))

    def _create_toolbar(self) -> None:
        toolbar = QToolBar("Formatting", self)
        toolbar.setMovable(False)
        toolbar.setFloatable(False)
        self.addToolBar(toolbar)

        # Create list dropdown (used by menu actions too)
        self._list_combo = QComboBox(self)
        self._list_combo.addItems([
            "None",
            "• Bullets",
            "1. Numbers (1, 2, 3)",
            "a. Letters (a, b, c)",
            "I. Roman (I, II, III)"
        ])
        self._list_combo.setCurrentIndex(0)
        self._list_combo.currentIndexChanged.connect(self._on_list_type_changed)

        # Typography
        self._font_combo = QFontComboBox(self)
        self._font_combo.setCurrentFont(QFont("Times New Roman"))
        self._font_combo.currentFontChanged.connect(self._on_font_family_changed)
        toolbar.addWidget(self._font_combo)

        self._size_combo = QComboBox(self)
        self._size_combo.setEditable(True)
        sizes = ["8", "9", "10", "11", "12", "14", "16", "18", "20", "22", "24", "26", "28", "36", "48", "72"]
        self._size_combo.addItems(sizes)
        self._size_combo.setCurrentText("12")
        self._size_combo.currentTextChanged.connect(self._on_font_size_changed)
        self._size_combo.setMaximumWidth(70)
        toolbar.addWidget(self._size_combo)

        toolbar.addSeparator()
        toolbar.addWidget(QLabel("List"))
        toolbar.addWidget(self._list_combo)


    def _insert_table(self) -> None:
        dlg = QDialog(self)
        dlg.setWindowTitle("Insert Table")
        dlg.setModal(True)

        rows = QSpinBox(dlg)
        rows.setMinimum(1)
        rows.setMaximum(50)
        rows.setValue(2)

        cols = QSpinBox(dlg)
        cols.setMinimum(1)
        cols.setMaximum(20)
        cols.setValue(2)

        form = QFormLayout()
        form.addRow("Rows:", rows)
        form.addRow("Columns:", cols)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(dlg.accept)
        buttons.rejected.connect(dlg.reject)

        layout = QVBoxLayout()
        layout.addLayout(form)
        layout.addWidget(buttons)
        dlg.setLayout(layout)

        if dlg.exec_() != QDialog.Accepted:
            return

        cursor = self._editor.textCursor()
        fmt = QTextTableFormat()
        fmt.setBorder(1)
        fmt.setBorderStyle(QTextTableFormat.BorderStyle_Solid)
        fmt.setCellPadding(6)
        fmt.setCellSpacing(0)
        fmt.setWidth(QTextLength(QTextLength.PercentageLength, 100))
        # Ensure borders are black
        fmt.setBorderBrush(QBrush(QColor("black")))
        cursor.insertTable(int(rows.value()), int(cols.value()), fmt)

    def _add_row_above(self) -> None:
        cursor = self._editor.textCursor()
        table = cursor.currentTable()
        if table:
            cell = table.cellAt(cursor)
            table.insertRows(cell.row(), 1)

    def _add_row_below(self) -> None:
        cursor = self._editor.textCursor()
        table = cursor.currentTable()
        if table:
            cell = table.cellAt(cursor)
            table.insertRows(cell.row() + 1, 1)

    def _remove_row(self) -> None:
        cursor = self._editor.textCursor()
        table = cursor.currentTable()
        if table:
            cell = table.cellAt(cursor)
            table.removeRows(cell.row(), 1)

    def _add_column_before(self) -> None:
        cursor = self._editor.textCursor()
        table = cursor.currentTable()
        if table:
            cell = table.cellAt(cursor)
            table.insertColumns(cell.column(), 1)

    def _add_column_after(self) -> None:
        cursor = self._editor.textCursor()
        table = cursor.currentTable()
        if table:
            cell = table.cellAt(cursor)
            table.insertColumns(cell.column() + 1, 1)

    def _remove_column(self) -> None:
        cursor = self._editor.textCursor()
        table = cursor.currentTable()
        if table:
            cell = table.cellAt(cursor)
            table.removeColumns(cell.column(), 1)

    def _resize_table(self) -> None:
        cursor = self._editor.textCursor()
        table = cursor.currentTable()
        if table is None:
            QMessageBox.information(self, "Table", "Place the cursor inside a table to resize it.")
            return

        dlg = QDialog(self)
        dlg.setWindowTitle("Table Size")
        dlg.setModal(True)

        width_pct = QSpinBox(dlg)
        width_pct.setMinimum(10)
        width_pct.setMaximum(100)
        width_pct.setValue(100)

        equal_cols = QCheckBox("Equal column widths", dlg)
        equal_cols.setChecked(True)

        form = QFormLayout()
        form.addRow("Table width (%):", width_pct)
        form.addRow("", equal_cols)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(dlg.accept)
        buttons.rejected.connect(dlg.reject)

        layout = QVBoxLayout()
        layout.addLayout(form)
        layout.addWidget(buttons)
        dlg.setLayout(layout)

        if dlg.exec_() != QDialog.Accepted:
            return

        tf = table.format()
        tf.setWidth(QTextLength(QTextLength.PercentageLength, float(width_pct.value())))

        if equal_cols.isChecked():
            cols = table.columns()
            if cols > 0:
                each = 100.0 / float(cols)
                tf.setColumnWidthConstraints(
                    [QTextLength(QTextLength.PercentageLength, each) for _ in range(cols)]
                )

        table.setFormat(tf)

    def _merge_char_format(self, fmt: QTextCharFormat) -> None:
        cursor = self._editor.textCursor()
        if not cursor.hasSelection():
            cursor.select(cursor.WordUnderCursor)
        cursor.mergeCharFormat(fmt)
        self._editor.mergeCurrentCharFormat(fmt)

    def _toggle_bold(self, checked: bool) -> None:
        fmt = QTextCharFormat()
        fmt.setFontWeight(QFont.Bold if checked else QFont.Normal)
        self._merge_char_format(fmt)

    def _toggle_italic(self, checked: bool) -> None:
        fmt = QTextCharFormat()
        fmt.setFontItalic(checked)
        self._merge_char_format(fmt)

    def _toggle_underline(self, checked: bool) -> None:
        fmt = QTextCharFormat()
        fmt.setFontUnderline(checked)
        self._merge_char_format(fmt)

    def _on_list_type_changed(self, index: int) -> None:
        cursor = self._editor.textCursor()
        
        # Remove any existing list formatting first
        if cursor.currentList():
            cursor.currentList().remove(cursor.block())
        
        # Apply new list formatting based on selection
        if index == 1:  # Bullets
            cursor.insertList(QTextListFormat.ListDisc)
        elif index == 2:  # Numbers (1, 2, 3)
            cursor.insertList(QTextListFormat.ListDecimal)
        elif index == 3:  # Letters (a, b, c)
            cursor.insertList(QTextListFormat.ListLowerAlpha)
        elif index == 4:  # Roman (I, II, III)
            cursor.insertList(QTextListFormat.ListUpperRoman)
        
        # Reset dropdown to "None" after applying
        if index != 0:
            self._list_combo.blockSignals(True)
            self._list_combo.setCurrentIndex(0)
            self._list_combo.blockSignals(False)

    def _on_font_family_changed(self, font: QFont) -> None:
        fmt = QTextCharFormat()
        fmt.setFontFamily(font.family())
        self._merge_char_format(fmt)

    def _on_font_size_changed(self, size_str: str) -> None:
        try:
            size = float(size_str)
            if size > 0:
                fmt = QTextCharFormat()
                fmt.setFontPointSize(size)
                self._merge_char_format(fmt)
        except ValueError:
            pass

    def _update_format_controls(self) -> None:
        fmt = self._editor.currentCharFormat()
        # Block signals so we don't trigger re-applying formats
        self._font_combo.blockSignals(True)
        self._font_combo.setCurrentFont(fmt.font())
        self._font_combo.blockSignals(False)

        self._size_combo.blockSignals(True)
        self._size_combo.setCurrentText(str(int(fmt.fontPointSize())) if fmt.fontPointSize() > 0 else "12")
        self._size_combo.blockSignals(False)

    def _window_title(self) -> str:
        base = self._doc.title if self._doc.title else "Untitled"
        if self._doc.id is None:
            return f"Signix - Editor ({base})"
        return f"Signix - Editor ({base}) [#{self._doc.id}]"

    def _update_status(self) -> None:
        state = "Unsaved" if self._dirty else "Saved"
        if self._doc.assigned_to is None:
            assigned = "None"
        else:
            u = self._workflow.get_user(self._doc.assigned_to)
            assigned = str(self._doc.assigned_to) if u is None else f"{u.username} ({u.role})"
        self.statusBar().showMessage(f"Status: {self._doc.status} | Assigned to: {assigned} | {state}")

    def _on_changed(self) -> None:
        if not self._dirty:
            self._dirty = True
        self._doc.title = self._title.text().strip() or "Untitled"
        self.setWindowTitle(self._window_title())
        self._update_status()

    def _sync_action_visibility(self) -> None:
        is_creator = self._doc.created_by == self._current_user.id or self._doc.id is None
        is_assignee = self._doc.assigned_to == self._current_user.id
        can_decide = is_assignee and self._doc.status == "Pending"
        
        # Creator can always edit; Approver can edit only when Pending and assigned to them.
        can_edit = is_creator or (is_assignee and self._doc.status == "Pending")
        
        self._title.setReadOnly(not can_edit)
        self._subject.setReadOnly(not can_edit)
        self._editor.setReadOnly(not can_edit)

        self._approve_btn.setEnabled(can_decide)
        self._reject_btn.setEnabled(can_decide)
        
        # Only creator can send for approval or set/change the approval chain
        # Initiator can send even if already Pending (to update/re-send the chain)
        self._send_btn.setEnabled(is_creator and self._doc.status in ("Draft", "Pending"))
        self._approval_chain_btn.setEnabled(is_creator)
        
        # Save button enabled if editing is allowed
        self._save_btn.setEnabled(can_edit)

    def _refresh_approval_chain_section(self) -> None:
        # Show configured approvers under the body (designation only)
        user_ids: List[int] = []
        if self._doc.id is None:
            user_ids = list(self._pending_chain_user_ids)
        else:
            try:
                chain = self._workflow.get_approval_chain(self._doc.id)
                user_ids = [int(s.user_id) for s in chain]
            except Exception:
                user_ids = []

        desigs: List[str] = []
        for uid in user_ids:
            u = self._workflow.get_user(int(uid))
            if u is None:
                continue
            desig = (getattr(u, "designation", "") or "").strip()
            if desig:
                desigs.append(desig)

        self._approval_chain_value.setText("\n\n".join(desigs))

    def configure_chain(self) -> None:
        initial = self._pending_chain_user_ids
        if self._doc.id is not None:
            chain = self._workflow.get_approval_chain(self._doc.id)
            initial = [s.user_id for s in chain]

        dlg = _ChainDialog(self._workflow, self._current_user, initial, self)
        if dlg.exec_() != QDialog.Accepted:
            return

        user_ids = dlg.user_ids()
        if not user_ids:
            QMessageBox.warning(self, "Chain", "Approval chain cannot be empty.")
            return

        if self._doc.id is None:
            self._pending_chain_user_ids = user_ids
            self._refresh_approval_chain_section()
            return

        try:
            self._workflow.set_approval_chain(document_id=self._doc.id, user_ids_in_order=user_ids)
        except Exception as exc:
            QMessageBox.critical(self, "Failed", str(exc))
            return

        self._refresh_approval_chain_section()

    def _sync_doc_from_ui(self) -> None:
        self._doc.title = self._title.text().strip() or "Untitled"
        self._doc.subject = self._subject.text().strip()
        self._doc.content = self._editor.toHtml()

    def save(self) -> None:
        # Log change if it's not a new doc AND (it's past Draft OR the editor is not the creator)
        has_id = self._doc.id is not None
        should_log_change = False
        if has_id and (self._doc.status != "Draft" or self._doc.created_by != self._current_user.id):
            try:
                old_doc = self._workflow.get_document(self._doc.id)
                # Simple HTML comparison (could be improved with a proper diff)
                if old_doc and (old_doc.content != self._editor.toHtml() or old_doc.title != self._title.text() or old_doc.subject != self._subject.text()):
                    should_log_change = True
            except Exception: pass

        self._sync_doc_from_ui()

        try:
            self._doc = self._workflow.save_document(self._doc)
            if should_log_change:
                self._comments.add_comment(
                    document_id=self._doc.id, 
                    user_id=self._current_user.id, 
                    comment=f"[System] {self._current_user.username} modified document content."
                )
                self._reload_comments()
        except Exception as exc:
            QMessageBox.critical(self, "Save failed", str(exc))
            return

        try:
            self._export_to_docx_template()
        except Exception as exc:
            # QMessageBox.warning(self, "Export failed", str(exc))
            pass

        if self._pending_chain_user_ids and self._doc.id is not None:
            try:
                self._workflow.set_approval_chain(
                    document_id=self._doc.id,
                    user_ids_in_order=list(self._pending_chain_user_ids),
                )
                self._pending_chain_user_ids = []
            except Exception as exc:
                QMessageBox.critical(self, "Failed", str(exc))
                return

        self._dirty = False
        self._sync_action_visibility()
        self.setWindowTitle(self._window_title())
        self._update_status()
        self._refresh_approval_chain_section()
        self._reload_comments()
        self.saved.emit()

    def send_for_approval(self) -> None:
        if self._doc.id is None or self._dirty:
            res = QMessageBox.question(
                self,
                "Save required",
                "Please save the document before sending for approval. Save now?",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.Yes,
            )
            if res == QMessageBox.Yes:
                self.save()
            else:
                return

        if self._doc.id is None:
            return

        # Capture initiator signature before sending.
        if not getattr(self._doc, "initiator_signature_png", None):
            sig_png = None
            try:
                sig_png = self._workflow.get_user_signature_png(user_id=self._current_user.id)
            except Exception:
                sig_png = None

            if not sig_png:
                dlg = _SignatureDialog(self)
                if dlg.exec_() != QDialog.Accepted:
                    return
                sig_png = dlg.signature_png_bytes()
                if not sig_png:
                    QMessageBox.warning(self, "Signature", "No signature provided. Send for approval cancelled.")
                    return
                try:
                    self._workflow.set_user_signature_png(user_id=self._current_user.id, signature_png=sig_png)
                except Exception:
                    pass
            self._doc.initiator_signature_png = sig_png
            try:
                self._doc = self._workflow.save_document(self._doc)
            except Exception as exc:
                QMessageBox.critical(self, "Failed", str(exc))
                return

        try:
            self._doc = self._workflow.send_for_approval(self._doc)
        except Exception as exc:
            QMessageBox.critical(self, "Failed", str(exc))
            return

        self._dirty = False
        self._sync_action_visibility()
        self._update_status()
        self._refresh_approval_chain_section()
        self.saved.emit()

    def approve(self) -> None:
        if self._doc.assigned_to != self._current_user.id:
            QMessageBox.warning(self, "Not allowed", "This document is not assigned to you.")
            return

        # First, ensure they have signed for this step
        res = QMessageBox.question(
            self,
            "Approve & Sign",
            "Are you sure you want to sign and approve this document?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.Yes,
        )
        if res != QMessageBox.Yes:
            return

        sig_png = None
        try:
            sig_png = self._workflow.get_user_signature_png(user_id=self._current_user.id)
        except Exception:
            sig_png = None

        if not sig_png:
            dlg = _SignatureDialog(self)
            if dlg.exec_() != QDialog.Accepted:
                return
            sig_png = dlg.signature_png_bytes()
            if not sig_png:
                QMessageBox.warning(self, "Sign", "No signature provided. Approval cancelled.")
                return
            try:
                self._workflow.set_user_signature_png(user_id=self._current_user.id, signature_png=sig_png)
            except Exception:
                pass

        try:
            # Save signature first
            self._workflow.set_approval_step_signature(
                document_id=int(self._doc.id),
                step_order=int(self._doc.current_step),
                signature_png=sig_png,
            )
            
            # Save current state (detecting if any edits were made)
            self.save()
            
            # Then approve
            self._doc = self._workflow.approve(self._doc)
        except Exception as exc:
            QMessageBox.critical(self, "Failed", str(exc))
            return

        self._sync_action_visibility()
        self._update_status()
        self._refresh_approval_chain_section()
        self.saved.emit()

    def reject(self) -> None:
        if self._doc.assigned_to != self._current_user.id:
            QMessageBox.warning(self, "Not allowed", "This document is not assigned to you.")
            return

        res = QMessageBox.question(
            self,
            "Reject",
            "Reject this document?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No,
        )
        if res != QMessageBox.Yes:
            return
        try:
            self._doc = self._workflow.reject(self._doc)
        except Exception as exc:
            QMessageBox.critical(self, "Failed", str(exc))
            return

        self._dirty = False
        self._sync_action_visibility()
        self._update_status()
        self._refresh_approval_chain_section()
        self.saved.emit()

    def _reload_comments(self) -> None:
        lines: List[str] = []
        if self._doc.id is None:
            self._comment_widget.set_comments(lines)
            self._comment_widget.set_add_enabled(False)
            return

        try:
            comments = self._comments.list_comments(self._doc.id)
        except Exception:
            self._comment_widget.set_comments(lines)
            return

        for c in comments:
            u = self._workflow.get_user(c.user_id)
            name = str(c.user_id) if u is None else u.username
            lines.append(f"{c.timestamp} - {name}: {c.comment}")

        self._comment_widget.set_comments(lines)
        self._comment_widget.set_add_enabled(True)

    def _on_comment_added(self, text: str) -> None:
        if self._doc.id is None:
            QMessageBox.warning(self, "Comments", "Save the document before adding comments.")
            return

        try:
            self._comments.add_comment(document_id=self._doc.id, user_id=self._current_user.id, comment=text)
        except Exception as exc:
            QMessageBox.critical(self, "Failed", str(exc))
            return

        self._reload_comments()
